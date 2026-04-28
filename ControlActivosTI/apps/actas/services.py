from decimal import Decimal
from io import BytesIO
from copy import copy
from math import ceil
from zipfile import ZIP_DEFLATED, ZipFile

from django.conf import settings
from django.core.files.base import ContentFile
from django.utils import timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as ExcelImage
from openpyxl.drawing.spreadsheet_drawing import AnchorMarker, OneCellAnchor
from openpyxl.drawing.xdr import XDRPositiveSize2D
from openpyxl.styles import Alignment

from .models import ActaEntrega


TIPO_ENTREGA = ActaEntrega.TipoActa.ENTREGA
TIPO_RECEPCION = ActaEntrega.TipoActa.RECEPCION
FILA_INICIO_ACTIVOS = 14
FILAS_ACTIVOS_PLANTILLA = 3
FILA_DESPUES_ACTIVOS = FILA_INICIO_ACTIVOS + FILAS_ACTIVOS_PLANTILLA
COLUMNAS_ACTIVOS = range(2, 10)
FORMATOS_LOGO = ("*.png", "*.jpg", "*.jpeg")
EMU_POR_PIXEL = 9525


def _texto(valor, default="-"):
    if valor is None:
        return default
    valor = str(valor).strip()
    return valor or default


def _moneda(valor):
    if valor in (None, ""):
        return "-"
    if isinstance(valor, Decimal):
        return f"USD {valor:,.2f}"
    try:
        return f"USD {Decimal(valor):,.2f}"
    except Exception:
        return _texto(valor)


def _nombre_archivo(asignacion, tipo, devolucion=None):
    codigo = asignacion.codigo_asignacion or f"ASG-{asignacion.pk}"
    if devolucion:
        codigo = devolucion.codigo_devolucion or f"{codigo}-DEV-{devolucion.pk}"
    prefijo = "acta_recepcion" if tipo == TIPO_RECEPCION else "acta_entrega"
    extension = "xlsx" if tipo == TIPO_ENTREGA and devolucion is None else "docx"
    return f"{prefijo}_{codigo}.{extension}"


def _fecha_acta(asignacion, tipo, devolucion=None):
    if devolucion:
        return devolucion.fecha_devolucion
    if tipo == TIPO_RECEPCION:
        return asignacion.fecha_devolucion or asignacion.fecha_asignacion
    return asignacion.fecha_asignacion


def _configuracion_acta(tipo):
    if tipo == TIPO_RECEPCION:
        return {
            "titulo": "ACTA DE RECEPCION DE BIENES INFORMATICOS",
            "intro": (
                "Por medio de la presente se deja constancia de la recepcion de los bienes "
                "informaticos detallados a continuacion, devueltos por el colaborador al area responsable."
            ),
            "declaracion": (
                "El area responsable declara haber recibido los bienes descritos y registra el estado "
                "final informado al momento de la devolucion."
            ),
            "firma_izquierda": "Entrega conforme",
            "firma_derecha": "Recibe",
        }
    return {
        "titulo": "ACTA DE ENTREGA DE BIENES INFORMATICOS",
        "intro": (
            "Por medio de la presente se deja constancia de la entrega de los bienes "
            "informaticos detallados a continuacion, los cuales quedan bajo custodia del colaborador."
        ),
        "declaracion": (
            "El colaborador declara haber recibido los bienes descritos en buen estado y se compromete "
            "a su uso adecuado, custodia y devolucion cuando corresponda."
        ),
        "firma_izquierda": "Entrega",
        "firma_derecha": "Recibe conforme",
    }


def _estado_detalle(detalle, tipo):
    if tipo == TIPO_RECEPCION and detalle.estado_activo_devolucion_id:
        return detalle.estado_activo_devolucion.nombre
    return detalle.activo.estado_activo.nombre


def _observaciones_detalle(detalle, tipo, devolucion_detalle=None):
    if devolucion_detalle:
        partes = []
        if devolucion_detalle.observaciones:
            partes.append(devolucion_detalle.observaciones.strip())
        if devolucion_detalle.devolucion.observaciones:
            partes.append(devolucion_detalle.devolucion.observaciones.strip())
        return " | ".join([parte for parte in partes if parte]) or "-"
    if tipo == TIPO_RECEPCION:
        partes = []
        if detalle.observaciones_devolucion:
            partes.append(detalle.observaciones_devolucion.strip())
        if detalle.asignacion.observaciones_devolucion:
            partes.append(detalle.asignacion.observaciones_devolucion.strip())
        return " | ".join([parte for parte in partes if parte]) or "-"
    return detalle.observaciones_acta


def _firmas(asignacion, tipo, devolucion=None):
    colaborador = f"{_texto(asignacion.nombre_colaborador_completo)}\n{_texto(asignacion.colaborador.cargo)}"
    responsable_entrega = _texto(
        asignacion.usuario_responsable.get_full_name() or asignacion.usuario_responsable.username
    )
    responsable_recepcion = responsable_entrega
    if devolucion:
        responsable_recepcion = _texto(
            devolucion.usuario_recepcion.get_full_name() or devolucion.usuario_recepcion.username
        )
    elif asignacion.usuario_recepcion_id:
        responsable_recepcion = _texto(
            asignacion.usuario_recepcion.get_full_name() or asignacion.usuario_recepcion.username
        )

    if tipo == TIPO_RECEPCION:
        return colaborador, responsable_recepcion
    return responsable_entrega, colaborador


def _plantilla_acta_entrega_path():
    plantilla_dir = settings.BASE_DIR / "templates" / "actas"
    plantillas = sorted(plantilla_dir.glob("*.xlsx"))
    if not plantillas:
        raise FileNotFoundError(f"No existe una plantilla .xlsx en {plantilla_dir}.")
    for plantilla in plantillas:
        nombre = plantilla.name.lower()
        if "entrega" in nombre or "asignaci" in nombre:
            return plantilla
    return plantillas[0]


def _logo_acta_path():
    plantilla_dir = settings.BASE_DIR / "templates" / "actas"
    for formato in FORMATOS_LOGO:
        for archivo in sorted(plantilla_dir.glob(formato)):
            if "logo_ilsa" in archivo.name.lower():
                return archivo
    for formato in FORMATOS_LOGO:
        for archivo in sorted(plantilla_dir.glob(formato)):
            if "logo" in archivo.name.lower():
                return archivo
    return None


def _valor_excel(valor):
    if valor in (None, ""):
        return None
    try:
        return float(Decimal(valor))
    except Exception:
        return _texto(valor, default="")


def _buscar_fila_por_texto(ws, texto):
    texto = texto.lower()
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is not None and texto in str(cell.value).lower():
                return cell.row
    return None


def _insertar_filas_preservando_combinadas(ws, indice, cantidad):
    if cantidad <= 0:
        return

    rangos = list(ws.merged_cells.ranges)
    for rango in rangos:
        ws.unmerge_cells(str(rango))

    ws.insert_rows(indice, cantidad)

    for rango in rangos:
        min_col, min_row, max_col, max_row = rango.bounds
        if min_row >= indice:
            min_row += cantidad
            max_row += cantidad
        elif max_row >= indice:
            max_row += cantidad
        ws.merge_cells(
            start_row=min_row,
            start_column=min_col,
            end_row=max_row,
            end_column=max_col,
        )


def _copiar_formato_fila(ws, fila_origen, fila_destino):
    ws.row_dimensions[fila_destino].height = ws.row_dimensions[fila_origen].height
    for columna in COLUMNAS_ACTIVOS:
        origen = ws.cell(fila_origen, columna)
        destino = ws.cell(fila_destino, columna)
        destino._style = copy(origen._style)
        destino.number_format = origen.number_format
        destino.font = copy(origen.font)
        destino.fill = copy(origen.fill)
        destino.border = copy(origen.border)
        destino.alignment = copy(origen.alignment)
        destino.protection = copy(origen.protection)
        destino.value = None


def _preparar_filas_activos(ws, cantidad_activos):
    filas_extra = max(0, cantidad_activos - FILAS_ACTIVOS_PLANTILLA)
    if filas_extra:
        _insertar_filas_preservando_combinadas(ws, FILA_DESPUES_ACTIVOS, filas_extra)
        for fila in range(FILA_DESPUES_ACTIVOS, FILA_DESPUES_ACTIVOS + filas_extra):
            _copiar_formato_fila(ws, FILA_DESPUES_ACTIVOS - 1, fila)
            ws.merge_cells(start_row=fila, start_column=2, end_row=fila, end_column=3)
            ws.merge_cells(start_row=fila, start_column=4, end_row=fila, end_column=5)


def _alineacion_con_ajuste(celda, vertical="top"):
    return Alignment(
        horizontal=celda.alignment.horizontal,
        vertical=vertical,
        text_rotation=celda.alignment.text_rotation,
        wrap_text=True,
        shrink_to_fit=celda.alignment.shrink_to_fit,
        indent=celda.alignment.indent,
    )


def _ajustar_alto_fila_activo(ws, fila):
    texto = _texto(ws.cell(fila, 8).value, default="")
    altura = 30.75
    if texto:
        altura = max(altura, min(95, 8 + ceil(len(texto) / 34) * 13.5))
    ws.row_dimensions[fila].height = altura
    for columna in COLUMNAS_ACTIVOS:
        celda = ws.cell(fila, columna)
        celda.alignment = _alineacion_con_ajuste(celda)


def _ajustar_bloques_largos(ws):
    inicio = _buscar_fila_por_texto(ws, "OBLIGACIONES Y DECLARACIONES")
    fin = _buscar_fila_por_texto(ws, "FOTOGRAFIAS DE LOS EQUIPOS") or _buscar_fila_por_texto(
        ws, "FOTOGRAFÍAS DE LOS EQUIPOS"
    )
    if not inicio or not fin:
        return

    for fila in range(inicio + 1, fin):
        celda = ws.cell(fila, 2)
        if not isinstance(celda.value, str) or len(celda.value.strip()) < 80:
            continue
        celda.alignment = _alineacion_con_ajuste(celda)
        ws.row_dimensions[fila].height = min(82, 5 + ceil(len(celda.value) / 155) * 13.5)


def _ancho_columnas_px(ws, columnas):
    ancho = 0
    for columna in columnas:
        width = ws.column_dimensions[columna].width or 8.43
        ancho += int(width * 7 + 5)
    return ancho


def _alto_filas_px(ws, filas):
    alto = 0
    for fila in filas:
        height = ws.row_dimensions[fila].height or 15
        alto += int(height * 96 / 72)
    return alto


def _sanear_texto_enriquecido_xlsx(contenido):
    entrada = BytesIO(contenido)
    salida = BytesIO()
    with ZipFile(entrada, "r") as origen, ZipFile(salida, "w", ZIP_DEFLATED) as destino:
        for item in origen.infolist():
            datos = origen.read(item.filename)
            if item.filename.startswith("xl/worksheets/") and item.filename.endswith(".xml"):
                datos = datos.replace(b"<t> </t>", b'<t xml:space="preserve"> </t>')
            destino.writestr(item, datos)
    salida.seek(0)
    return salida.getvalue()


def _colocar_logo(ws):
    ws["B1"] = None
    logo_path = _logo_acta_path()
    if not logo_path:
        return

    logo = ExcelImage(logo_path)
    max_width = _ancho_columnas_px(ws, ("B", "C")) - 8
    max_height = _alto_filas_px(ws, range(1, 6)) - 6
    escala = min(max_width / logo.width, max_height / logo.height, 1)
    logo.width = int(logo.width * escala)
    logo.height = int(logo.height * escala)
    offset_x = max(0, int((max_width - logo.width) / 2) + 4)
    offset_y = max(0, int((max_height - logo.height) / 2) + 3)
    logo.anchor = OneCellAnchor(
        _from=AnchorMarker(
            col=1,
            row=0,
            colOff=offset_x * EMU_POR_PIXEL,
            rowOff=offset_y * EMU_POR_PIXEL,
        ),
        ext=XDRPositiveSize2D(
            cx=logo.width * EMU_POR_PIXEL,
            cy=logo.height * EMU_POR_PIXEL,
        ),
    )
    ws.add_image(logo)


def _llenar_firmas_entrega(ws, asignacion):
    fila_recibe = _buscar_fila_por_texto(ws, "RECIBE CONFORME")
    if fila_recibe:
        ws.cell(fila_recibe + 1, 4).value = _texto(asignacion.nombre_colaborador_completo, default="")
        ws.cell(fila_recibe + 2, 4).value = _texto(asignacion.colaborador.cargo, default="")


def construir_hoja_acta_entrega(asignacion):
    workbook = load_workbook(_plantilla_acta_entrega_path(), rich_text=True)
    ws = workbook.active

    detalles = list(
        asignacion.detalles.select_related(
            "activo__tipo_activo",
            "activo__estado_activo",
        ).order_by("orden", "id")
    )
    _preparar_filas_activos(ws, len(detalles))

    ws["E6"] = timezone.localdate()
    ws["E10"] = _texto(asignacion.nombre_colaborador_completo, default="")
    ws["E11"] = _texto(asignacion.colaborador.cedula, default="")
    ws["I10"] = _texto(asignacion.colaborador.cargo, default="")
    _colocar_logo(ws)

    for indice, detalle in enumerate(detalles, start=FILA_INICIO_ACTIVOS):
        ws.cell(indice, 2).value = _texto(detalle.articulo_acta, default="")
        ws.cell(indice, 4).value = _texto(detalle.activo.marca, default="")
        ws.cell(indice, 6).value = _valor_excel(detalle.activo.valor)
        ws.cell(indice, 6).number_format = '$#,##0.00'
        ws.cell(indice, 7).value = None
        ws.cell(indice, 8).value = detalle.caracteristicas_acta if detalle.caracteristicas_acta != "-" else ""
        ws.cell(indice, 9).value = None
        _ajustar_alto_fila_activo(ws, indice)

    _llenar_firmas_entrega(ws, asignacion)

    salida = BytesIO()
    workbook.save(salida)
    salida.seek(0)
    return _sanear_texto_enriquecido_xlsx(salida.getvalue())


def construir_documento_acta(asignacion, tipo=TIPO_ENTREGA, devolucion=None):
    config = _configuracion_acta(tipo)
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    titulo = document.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(config["titulo"])
    run.bold = True
    run.font.size = Pt(14)

    subtitulo = document.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if devolucion:
        subtitulo.add_run(f"Codigo de devolucion: {_texto(devolucion.codigo_devolucion)}").bold = True
        referencia = document.add_paragraph()
        referencia.alignment = WD_ALIGN_PARAGRAPH.CENTER
        referencia.add_run(f"Asignacion original: {_texto(asignacion.codigo_asignacion)}")
    else:
        subtitulo.add_run(f"Codigo de asignacion: {_texto(asignacion.codigo_asignacion)}").bold = True

    document.add_paragraph()

    fecha_acta = _fecha_acta(asignacion, tipo, devolucion=devolucion)
    datos = [
        ("Fecha de suscripcion", fecha_acta.strftime("%d/%m/%Y")),
        ("Colaborador", asignacion.nombre_colaborador_completo),
        ("Cedula", _texto(asignacion.colaborador.cedula)),
        ("Cargo", _texto(asignacion.colaborador.cargo)),
        ("Area", _texto(asignacion.colaborador.area)),
        ("Empresa", _texto(asignacion.colaborador.empresa)),
        ("Ubicacion", _texto(asignacion.colaborador.ubicacion)),
    ]

    tabla_datos = document.add_table(rows=0, cols=2)
    tabla_datos.style = "Table Grid"
    for etiqueta, valor in datos:
        row = tabla_datos.add_row().cells
        row[0].text = etiqueta
        row[1].text = _texto(valor)

    document.add_paragraph()
    document.add_paragraph(config["intro"])

    document.add_paragraph()

    tabla = document.add_table(rows=1, cols=6)
    tabla.style = "Table Grid"
    encabezados = ["Articulo", "Marca", "Valor", "Estado", "Caracteristicas", "Observaciones"]
    for idx, texto in enumerate(encabezados):
        tabla.rows[0].cells[idx].text = texto

    if devolucion:
        detalles = devolucion.detalles.select_related(
            "detalle_asignacion__activo__tipo_activo",
            "detalle_asignacion__activo__estado_activo",
            "estado_activo_devolucion",
            "devolucion",
        ).order_by("detalle_asignacion__orden", "id")
    else:
        detalles = asignacion.detalles.select_related(
            "activo__tipo_activo",
            "activo__estado_activo",
            "estado_activo_devolucion",
        ).order_by("orden", "id")

    for item in detalles:
        detalle = item.detalle_asignacion if devolucion else item
        row = tabla.add_row().cells
        row[0].text = _texto(detalle.articulo_acta)
        row[1].text = _texto(detalle.activo.marca)
        row[2].text = _moneda(detalle.activo.valor)
        estado = item.estado_activo_devolucion.nombre if devolucion else _estado_detalle(detalle, tipo)
        row[3].text = _texto(estado)
        row[4].text = detalle.caracteristicas_acta
        row[5].text = _observaciones_detalle(detalle, tipo, devolucion_detalle=item if devolucion else None)

    document.add_paragraph()
    document.add_paragraph(config["declaracion"])

    document.add_paragraph()
    firmas = document.add_table(rows=2, cols=2)
    firmas.style = "Table Grid"
    firmas.cell(0, 0).text = config["firma_izquierda"]
    firmas.cell(0, 1).text = config["firma_derecha"]
    firma_izquierda, firma_derecha = _firmas(asignacion, tipo, devolucion=devolucion)
    firmas.cell(1, 0).text = firma_izquierda
    firmas.cell(1, 1).text = firma_derecha

    salida = BytesIO()
    document.save(salida)
    salida.seek(0)
    return salida.getvalue()


def generar_o_actualizar_acta(asignacion, usuario, tipo=TIPO_ENTREGA, devolucion=None):
    if tipo == TIPO_ENTREGA and devolucion is None:
        contenido = construir_hoja_acta_entrega(asignacion)
    else:
        contenido = construir_documento_acta(asignacion, tipo=tipo, devolucion=devolucion)
    nombre_archivo = _nombre_archivo(asignacion, tipo, devolucion=devolucion)

    filtros = {"asignacion": asignacion, "tipo": tipo, "devolucion": devolucion}
    acta, _ = ActaEntrega.objects.get_or_create(
        **filtros,
        defaults={"usuario_generador": usuario},
    )
    acta.usuario_generador = usuario
    acta.nombre_archivo = nombre_archivo
    acta.archivo.save(nombre_archivo, ContentFile(contenido), save=False)
    acta.save()
    return acta


def generar_o_actualizar_actas_devolucion(devolucion, usuario):
    asignacion = devolucion.asignacion
    acta_entrega = ActaEntrega.objects.filter(
        asignacion=asignacion,
        tipo=TIPO_ENTREGA,
        devolucion__isnull=True,
        archivo__isnull=False,
    ).exclude(archivo="").first()
    if not acta_entrega:
        acta_entrega = generar_o_actualizar_acta(asignacion, usuario, tipo=TIPO_ENTREGA)
    acta_recepcion = generar_o_actualizar_acta(
        asignacion,
        usuario,
        tipo=TIPO_RECEPCION,
        devolucion=devolucion,
    )
    return acta_entrega, acta_recepcion
